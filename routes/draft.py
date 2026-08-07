from flask import Blueprint, jsonify, request, send_file
from flask_login import current_user, login_required
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.context import get_context as get_stored_context
from services.jobs import create_job, update_job
from services.request_context import resolve_matter_id
from services.task_queue import enqueue_job
from services.tenancy import AuthorizationError


draft_bp = Blueprint("draft", __name__)


@draft_bp.route("/draft", methods=["POST"])
@login_required
def draft():
    """Queue a matter_draft job; document generation runs in the job."""
    payload = request.json or {}
    context_id = resolve_matter_id(payload)
    doc_type = str(payload.get("doc_type") or "memo")  # "memo" or "brief"

    if not context_id:
        return jsonify({"error": "No context_id provided"}), 400

    uid = str(current_user.get_id())
    try:
        job, created = create_job(context_id, uid, "matter_draft", {"doc_type": doc_type})
        if created:
            enqueue_job(context_id, job["job_id"])
    except AuthorizationError:
        return jsonify({"error": "forbidden"}), 403
    except Exception as exc:
        if "job" in locals():
            update_job(context_id, job["job_id"], status="failed", stage="enqueue_failed",
                       error={"code": "enqueue_failed", "message": str(exc)[:300]})
        return jsonify({"error": "unable to queue draft"}), 503
    job["status_url"] = f"/api/matters/{context_id}/jobs/{job['job_id']}"
    job["context_id"] = context_id
    return jsonify(job), 202


@draft_bp.route("/draft/export", methods=["POST"])
@login_required
def draft_export():
    payload = request.json or {}
    context_id = resolve_matter_id(payload)
    if not context_id:
        return jsonify({"error": "No context_id provided"}), 400

    context = get_stored_context(context_id, str(current_user.id))
    if not context:
        return jsonify({"error": "Context not found"}), 404

    draft_text = payload.get("draft_text") or context.get("draft")
    if not draft_text:
        return jsonify({"error": "No draft found"}), 404

    # Remove formatting artifacts
    draft_text = re.sub(r'[*#]', '', draft_text)

    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('MEMORANDUM')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    case_title = context.get('title', 'Untitled Case')
    
    info_lines = [
        ('TO:', '[Recipient]'),
        ('FROM:', '[Author]'),
        ('RE:', case_title),
        ('DATE:', __import__('datetime').datetime.now().strftime('%B %d, %Y'))
    ]
    
    for label, value in info_lines:
        p = doc.add_paragraph()
        run_label = p.add_run(label + '\t')
        run_label.bold = True
        p.add_run(value)
    
    divider = doc.add_paragraph('_' * 70)
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = draft_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        if line.isupper() and len(line) > 3:
            heading = doc.add_paragraph()
            heading_run = heading.add_run(line)
            heading_run.bold = True
            heading_run.font.size = Pt(13)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"{context.get('title', 'legal_memo').replace(' ', '_')}.docx"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
