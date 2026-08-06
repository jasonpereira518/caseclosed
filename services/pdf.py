import os
import uuid
import logging
import io
from docx import Document as DocxDocument
from pdfminer.high_level import extract_text as extract_pdf_text_miner
from werkzeug.utils import secure_filename

def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx', 'txt'}

def secure_save_document(file_obj, upload_folder: str) -> tuple[str, str]:
    """Generates a UUID-locked filename preventing cross-request race conditions."""
    safe_name = secure_filename(file_obj.filename)
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"
    path = os.path.join(upload_folder, unique_name)
    file_obj.save(path)
    return safe_name, path

def extract_pdf_text(filepath: str) -> str:
    try:
        pdf_text = extract_pdf_text_miner(filepath)
        if not pdf_text or len(pdf_text.strip()) < 40:
            return extract_pdf_with_document_ai(filepath)
        return pdf_text
    except Exception as e:
        try:
            return extract_pdf_with_document_ai(filepath)
        except Exception as ocr_error:
            raise ValueError(f"PDF extraction failed: {ocr_error}") from e


def extract_pdf_with_document_ai(filepath: str) -> str:
    """OCR a PDF through Document AI when a processor is configured."""
    import config
    if not config.DOCUMENT_AI_PROCESSOR_ID or not config.PROJECT_ID:
        raise ValueError("PDF contains too little embedded text and Document AI OCR is not configured")
    from google.cloud import documentai
    client = documentai.DocumentProcessorServiceClient(
        client_options={"api_endpoint": f"{config.DOCUMENT_AI_LOCATION}-documentai.googleapis.com"})
    name = client.processor_path(config.PROJECT_ID, config.DOCUMENT_AI_LOCATION,
                                 config.DOCUMENT_AI_PROCESSOR_ID)
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(filepath)
    texts = []
    # Online OCR processors accept bounded page counts. Split in memory so the
    # original never needs a second durable copy.
    for start in range(0, len(reader.pages), 15):
        writer = PdfWriter()
        for page in reader.pages[start:start + 15]:
            writer.add_page(page)
        buffer = io.BytesIO()
        writer.write(buffer)
        response = client.process_document(request=documentai.ProcessRequest(
            name=name, raw_document=documentai.RawDocument(
                content=buffer.getvalue(), mime_type="application/pdf")))
        texts.append(str(response.document.text or "").strip())
    text = "\n\n".join(part for part in texts if part).strip()
    if not text:
        raise ValueError("Document AI returned no text")
    return text

def extract_docx_text(filepath: str) -> str:
    try:
        doc = DocxDocument(filepath)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}") from e

def extract_txt_text(filepath: str) -> str:
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def extract_document_text(filepath: str, original_filename: str) -> str:
    ext = original_filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        return extract_pdf_text(filepath)
    elif ext == 'docx':
        return extract_docx_text(filepath)
    elif ext == 'doc':
        raise ValueError("Legacy .doc files are not supported; save the file as .docx or PDF")
    elif ext == 'txt':
        return extract_txt_text(filepath)
    else:
        raise ValueError(f'Unsupported file type: .{ext}')

def cleanup_temp_file(filepath: str):
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
    except Exception as e:
        logging.error("Failed to delete temp file %s: %s", filepath, e)
